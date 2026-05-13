"""Data Analysis Agent — supports CSV, Excel, JSON, TSV, Parquet, DOCX."""
from __future__ import annotations
import logging
import pandas as pd
from core.schemas import ColumnStat, DatasetMetadata
from core.llm import chat

logger = logging.getLogger(__name__)


def _load_dataframe(filepath: str) -> pd.DataFrame:
    fp = filepath.lower()
    try:
        if fp.endswith((".xlsx", ".xls", ".xlsm", ".xlsb")):
            return pd.read_excel(filepath)
        elif fp.endswith(".json"):
            try:
                return pd.read_json(filepath)
            except Exception:
                import json
                with open(filepath) as f:
                    data = json.load(f)
                if isinstance(data, list):
                    return pd.DataFrame(data)
                elif isinstance(data, dict):
                    return pd.DataFrame([data])
                return pd.DataFrame()
        elif fp.endswith(".tsv"):
            return pd.read_csv(filepath, sep="\t", encoding_errors="replace")
        elif fp.endswith(".parquet"):
            return pd.read_parquet(filepath)
        elif fp.endswith(".docx"):
            return _read_docx(filepath)
        else:
            # Try CSV with multiple encodings
            for enc in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return pd.read_csv(filepath, encoding=enc, encoding_errors="replace")
                except Exception:
                    continue
            return pd.read_csv(filepath, encoding_errors="replace")
    except Exception as e:
        logger.error(f"Failed to load {filepath}: {e}")
        raise


def _read_docx(filepath: str) -> pd.DataFrame:
    try:
        from docx import Document
        doc = Document(filepath)
        tables = []
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
            tables.append(pd.DataFrame(rows, columns=headers))
        if tables:
            return pd.concat(tables, ignore_index=True)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return pd.DataFrame({"text": paragraphs})
    except ImportError:
        raise RuntimeError("python-docx not installed.")


def _compute_column_stats(df: pd.DataFrame) -> list[ColumnStat]:
    stats = []
    for col in df.columns:
        series = df[col]
        try:
            sample = series.dropna().head(5).tolist()
            # Convert non-serializable types
            sample = [str(v) if not isinstance(v, (int, float, str, bool)) else v for v in sample]
            stat = ColumnStat(
                name=str(col),
                dtype=str(series.dtype),
                non_null_count=int(series.notna().sum()),
                null_count=int(series.isna().sum()),
                unique_count=int(series.nunique()),
                sample_values=sample,
            )
            if pd.api.types.is_numeric_dtype(series) and series.notna().sum() > 0:
                stat.mean   = round(float(series.mean()), 4)
                stat.std    = round(float(series.std()),  4)
                stat.min    = round(float(series.min()),  4)
                stat.max    = round(float(series.max()),  4)
                stat.median = round(float(series.median()), 4)
            stats.append(stat)
        except Exception as e:
            logger.warning(f"Could not compute stats for column '{col}': {e}")
    return stats


def _compute_correlations(df: pd.DataFrame) -> dict[str, dict[str, float]]:
    try:
        num_df = df.select_dtypes(include="number")
        if num_df.shape[1] < 2:
            return {}
        corr = num_df.corr().round(3)
        return {str(col): {str(k): float(v) for k, v in corr[col].items()}
                for col in corr.columns}
    except Exception as e:
        logger.warning(f"Correlation failed: {e}")
        return {}


def _detect_anomalies(df: pd.DataFrame) -> list[str]:
    anomalies = []
    try:
        for col in df.select_dtypes(include="number").columns:
            series = df[col].dropna()
            if len(series) < 4:
                continue
            q1, q3 = series.quantile(0.25), series.quantile(0.75)
            iqr = q3 - q1
            if iqr == 0:
                continue
            outliers = series[(series < q1 - 3 * iqr) | (series > q3 + 3 * iqr)]
            if not outliers.empty:
                anomalies.append(f"Column '{col}' has {len(outliers)} extreme outlier(s).")
        missing_pct = (df.isna().sum() / len(df) * 100).round(1)
        for col, pct in missing_pct.items():
            if pct > 20:
                anomalies.append(f"Column '{col}' is {pct:.0f}% missing.")
    except Exception as e:
        logger.warning(f"Anomaly detection failed: {e}")
    return anomalies[:10]


async def run(filepath: str) -> DatasetMetadata:
    df = _load_dataframe(filepath)
    df.columns = [str(c).strip() for c in df.columns]
    # Drop completely empty columns
    df = df.dropna(axis=1, how="all")
    df_sample = df.head(5000)

    columns      = _compute_column_stats(df_sample)
    correlations = _compute_correlations(df_sample)
    anomalies    = _detect_anomalies(df_sample)

    col_names = [c.name for c in columns][:20]
    domain = "general"
    try:
        prompt = (
            f"Given these dataset column names: {col_names}, "
            "in one word, what business domain is this? "
            "Examples: sales, finance, healthcare, marketing, logistics, hr, ecommerce. "
            "Reply with only the single domain word."
        )
        domain = chat(prompt).strip().lower().split()[0]
        # Sanitize — only allow known domains
        known = {"sales","finance","healthcare","marketing","logistics","hr","ecommerce",
                 "general","retail","education","operations","supply","inventory"}
        if domain not in known:
            domain = "general"
    except Exception as e:
        logger.warning(f"Domain inference failed: {e}")

    logger.info(f"Analysis: {len(df)} rows, {len(df.columns)} cols, domain={domain}")
    return DatasetMetadata(
        row_count=len(df), column_count=len(df.columns),
        columns=columns, correlations=correlations,
        anomalies=anomalies, inferred_domain=domain,
    )
